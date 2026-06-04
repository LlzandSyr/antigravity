import docx
import os

p2_path = r"c:\Users\86135\Desktop\面试题\阶段项目2\FluxCloud_V2.0_项目功能逐条复盘（含AI扩展）.docx"
p3_path = r"c:\Users\86135\Desktop\面试题\阶段项目3\Aura_V2.0_项目功能逐条复盘（含AI扩展）.docx"

def check_docx(path):
    if not os.path.exists(path):
        print(f"File {os.path.basename(path)} not found")
        return
    print(f"\n========================================\nChecking file: {os.path.basename(path)}\n========================================")
    doc = docx.Document(path)
    
    # Check paragraphs
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if any(marker in text for marker in ["TODO", "todo", "待补充", "待填写", "【", "暂无"]):
            print(f"Paragraph {i}: {text}")
            
    # Check tables
    for i, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if any(marker in text for marker in ["TODO", "todo", "待补充", "待填写", "【", "暂无"]):
                    print(f"Table {i}, Row {r_idx}, Col {c_idx}: {text[:100]}")

check_docx(p2_path)
check_docx(p3_path)
