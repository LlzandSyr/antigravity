import docx
import os

p2_path = r"c:\Users\86135\Desktop\面试题\阶段项目2\FluxCloud_V2.0_项目功能逐条复盘（含AI扩展）.docx"
p3_path = r"c:\Users\86135\Desktop\面试题\阶段项目3\Aura_V2.0_项目功能逐条复盘（含AI扩展）.docx"

def replace_in_p(p, old, new):
    replaced = False
    if old in p.text:
        # Try to replace in runs to preserve formatting
        for run in p.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                replaced = True
        
        # If it was split across runs, fallback to paragraph level replacement
        if old in p.text:
            p.text = p.text.replace(old, new)
            replaced = True
    return replaced

def update_document(path, replacements):
    if not os.path.exists(path):
        print(f"File {os.path.basename(path)} not found")
        return
    
    print(f"\nUpdating {os.path.basename(path)}...")
    doc = docx.Document(path)
    
    counts = {old: 0 for old, _ in replacements}
    
    # Process paragraphs
    for p in doc.paragraphs:
        for old, new in replacements:
            if replace_in_p(p, old, new):
                counts[old] += 1
                
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in replacements:
                        if replace_in_p(p, old, new):
                            counts[old] += 1
                            
    for old, count in counts.items():
        print(f"  Replaced '{old}' -> {count} times")
        
    doc.save(path)
    print(f"Saved {os.path.basename(path)} successfully.")

# Phase 2 Replacements
p2_reps = [
    ("dir_file_list/dir_file_list.c（1167行）", "dir_file_list/dir_file_list.c（阶段二最终版本已扩展至2684行，以下已适配实际行号）"),
    ("第489~693行", "第916~1128行"),
    ("第787~803行", "第1222~1238行"),
    ("第941~960行", "第1376~1395行"),
    ("第809~935行", "第1244~1370行"),
    ("第453~482行", "第880~909行"),
    ("第403~447行", "第830~874行"),
    ("第199~384行", "第564~811行"),
    ("第1124~1166行", "第1559~1603行"),
    ("第741~781行", "第1176~1216行"),
    ("第968~1087行", "第1403~1522行"),
    ("第1093~1107行", "第1528~1542行"),
    ("第2551行", "第2557行"),
    ("第2568-2578行", "第2564~2586行"),
    ("第2609行", "第2614行"),
    ("第2638-2670行", "第2639~2670行")
]

# Phase 3 Replacements
p3_reps = [
    ("changeQty(): menudelegate.cpp 第209~220行", "changeQty(): menumodel.cpp 第77~90行"),
    ("Aura_Client/menumodel.cpp 第45~70行", "Aura_Client/menumodel.cpp 第49~68行"),
    ("Aura_Client/clientwindow.cpp 第364~386行", "Aura_Client/clientwindow.cpp 第339~412行"),
    ("menudelegate.cpp 第209~220行", "menumodel.cpp 第77~90行"), # fallback just in case
]

update_document(p2_path, p2_reps)
update_document(p3_path, p3_reps)
