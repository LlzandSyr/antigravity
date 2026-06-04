import docx
import os
import shutil

p2_v2 = r"c:\Users\86135\Desktop\面试题\阶段项目2\FluxCloud_V2.0_项目功能逐条复盘（含AI扩展）.docx"
p3_v2 = r"c:\Users\86135\Desktop\面试题\阶段项目3\Aura_V2.0_项目功能逐条复盘（含AI扩展）.docx"

p2_v3 = r"c:\Users\86135\Desktop\面试题\阶段项目2\FluxCloud_V3.0_项目功能逐条复盘（含AI扩展）.docx"
p3_v3 = r"c:\Users\86135\Desktop\面试题\阶段项目3\Aura_V3.0_项目功能逐条复盘（含AI扩展）.docx"

def replace_in_p(p, old, new):
    replaced = False
    if old in p.text:
        for run in p.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                replaced = True
        if old in p.text:
            p.text = p.text.replace(old, new)
            replaced = True
    return replaced

def apply_replacements(doc, reps):
    for p in doc.paragraphs:
        for old, new in reps:
            replace_in_p(p, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in reps:
                        replace_in_p(p, old, new)

# 1. Restore V2.0 documents
print("Restoring V2.0 documents to original state...")

p2_restore_reps = [
    ("dir_file_list/dir_file_list.c（阶段二最终版本已扩展至2684行，以下已适配实际行号）", "dir_file_list/dir_file_list.c（1167行）"),
    ("第916~1128行", "第489~693行"),
    ("第1222~1238行", "第787~803行"),
    ("第1376~1395行", "第941~960行"),
    ("第1244~1370行", "第809~935行"),
    ("第880~909行", "第453~482行"),
    ("第830~874行", "第403~447行"),
    ("第564~811行", "第199~384行"),
    ("第1559~1603行", "第1124~1166行"),
    ("第1176~1216行", "第741~781行"),
    ("第1403~1522行", "第968~1087行"),
    ("第1528~1542行", "第1093~1107行"),
    ("第2557行", "第2551行"),
    ("第2564~2586行", "第2568-2578行"),
    ("第2614行", "第2609行"),
    ("第2639~2670行", "第2638-2670行")
]

p3_restore_reps = [
    ("changeQty(): menumodel.cpp 第77~90行", "changeQty(): menudelegate.cpp 第209~220行"),
    ("loadFromJson(): 第49~68行", "loadFromJson(): 第45~70行"),
    ("OrderPage::initUI(): Aura_Client/clientwindow.cpp 第339~412行", "OrderPage::initUI(): Aura_Client/clientwindow.cpp 第364~386行"),
]

if os.path.exists(p2_v2):
    doc2 = docx.Document(p2_v2)
    apply_replacements(doc2, p2_restore_reps)
    doc2.save(p2_v2)
    print("Restored FluxCloud V2.0.")

if os.path.exists(p3_v2):
    doc3 = docx.Document(p3_v2)
    apply_replacements(doc3, p3_restore_reps)
    doc3.save(p3_v2)
    print("Restored Aura V2.0.")

# 2. Copy V2.0 to V3.0 documents
print("Copying V2.0 to V3.0...")
shutil.copy2(p2_v2, p2_v3)
shutil.copy2(p3_v2, p3_v3)

# 3. Apply updates to V3.0 documents
print("Applying updates to V3.0 documents...")

p2_v3_reps = [
    ("FluxCloud 流光云盘（阶段二）", "FluxCloud 流光云盘（阶段三）"), # Update version strings in text/headers
    ("FluxCloud_V2.0", "FluxCloud_V3.0"),
    ("dir_file_list/dir_file_list.c（1167行）", "dir_file_list/dir_file_list.c（在阶段三中已扩展至2684行，以下已适配实际行号）"),
    ("第489~693行", "第916~1128行"),
    ("第787~803行", "第1222~1238行"),
    ("第941~960行", "第1376~1395行"),
    ("第809~935行", "第1244~1370行"),
    ("第453~482行", "第880~909行"),
    ("第403~447行", "第830~874行"),
    ("第199~384行", "第564~811行"),
    ("第1124~1166行", "第1559~1603行"),
    ("第741~781行", "第1176~1216行"),
    ("第968~1087行", "第1403~1522行"),
    ("第1093~1107行", "第1528~1542行"),
    ("第2551行", "第2557行"),
    ("第2568-2578行", "第2564~2586行"),
    ("第2609行", "第2614行"),
    ("第2638-2670行", "第2639~2670行")
]

p3_v3_reps = [
    ("Aura_V2.0", "Aura_V3.0"),
    ("changeQty(): menudelegate.cpp 第209~220行", "changeQty(): menumodel.cpp 第77~90行"),
    ("loadFromJson(): 第45~70行", "loadFromJson(): 第49~68行"),
    ("OrderPage::initUI(): Aura_Client/clientwindow.cpp 第364~386行", "OrderPage::initUI(): Aura_Client/clientwindow.cpp 第339~412行"),
]

# Save V3.0 documents
doc2_v3 = docx.Document(p2_v3)
apply_replacements(doc2_v3, p2_v3_reps)
doc2_v3.save(p2_v3)
print("Updated and saved FluxCloud V3.0.")

doc3_v3 = docx.Document(p3_v3)
apply_replacements(doc3_v3, p3_v3_reps)
doc3_v3.save(p3_v3)
print("Updated and saved Aura V3.0.")

print("All tasks completed successfully!")
