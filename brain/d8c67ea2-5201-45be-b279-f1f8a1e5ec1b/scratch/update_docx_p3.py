import docx
import os

p3_path = r"c:\Users\86135\Desktop\面试题\阶段项目3\Aura_V2.0_项目功能逐条复盘（含AI扩展）.docx"

def replace_in_p(p, old, new):
    replaced = False
    if old in p.text:
        for run in p.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                replaced = True
        if old in p.text:
            p.text = p.text.replace(old, new)
            replaced = True
    return replaced

def update_document(path, replacements):
    if not os.path.exists(path):
        print(f"File {os.path.basename(path)} not found")
        return
    
    doc = docx.Document(path)
    counts = {old: 0 for old, _ in replacements}
    
    for p in doc.paragraphs:
        for old, new in replacements:
            if replace_in_p(p, old, new):
                counts[old] += 1
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in replacements:
                        if replace_in_p(p, old, new):
                            counts[old] += 1
                            
    for old, count in counts.items():
        print(f"  Replaced '{old}' -> {count} times")
        
    doc.save(path)
    print(f"Saved {os.path.basename(path)} successfully.")

# Remaining Aura Replacements
p3_reps = [
    ("loadFromJson(): 第45~70行", "loadFromJson(): 第49~68行"),
    ("menudelegate.cpp 第1~220行", "menudelegate.cpp 第1~208行"),
]

update_document(p3_path, p3_reps)
